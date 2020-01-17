import docx
import PyPDF2
from tabula import read_pdf
import tabula

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine


def getTable():
    doc = docx.Document('../dicts/word_test.docx')
    tables = doc.tables
    for table in tables:
        for row in table.rows:
            print()
            temp = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    temp.append(paragraph.text)
            print(temp)


def getTablepdf():
    # df = read_pdf('C:/Users/wenji/Desktop/subway.pdf', multiple_tables=True)
    tabula.convert_into('C:/Users/wenji/Desktop/subway.pdf', "C:/Users/wenji/Desktop/result/output.csv",output_format="csv", pages="all",  multiple_tables=True )
    # return df




def getPdf(filename):

    fp = open(filename, 'rb')
    parser = PDFParser(fp)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # Process each page contained in the document.
    result=[]
    for page in doc.get_pages():
        interpreter.process_page(page)
        layout = device.get_result()
        for lt_obj in layout:
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                result.append(lt_obj.get_text())
                # print(lt_obj.get_text())
    return result


def getTextPDF(filename):
    file  = open(filename, 'rb')
    reader= PyPDF2.PdfFileReader(file)
    page= reader.getPage(0).extractText()
    return page


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        if para.text != '':
           fullText.append(para.text)
    return fullText


def getBold(filename):
    doc = docx.Document(filename)

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and run.text != " ":
                print(run.text)




if __name__ == '__main__':


    r= getTablepdf ()
    # data is a dictionary

    # data = r[2]
    # print(data)
    # print(data.shape)
    # data.to_excel("C:/Users/wenji/Desktop/result/output.xlsx")


    # a= getBold('../dicts/test.docx')
    # import re
    # # b = getPdf('Larry.pdf')
    # # b = ','.join(b)
    # b = '*(&*^^&^%$Y%fhjdkhfjkdhf763662'
    # s = re.sub("\d+", '', b)
    # print(s)
    # s = re.sub("[^a-zA-Z0-9]+", ' ', s)
    # print(s)


    # text = getText('../dicts/word_test.docx')
    # for item in text:
    #     print(item)














